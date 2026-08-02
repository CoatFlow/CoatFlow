"""
Tests voor sjabloon_templatiseer(): een veld waarvan de (eventueel handmatig
gecorrigeerde) tekst NERGENS letterlijk in het geüploade voorbeelddocument
voorkomt, mag niet meer stil "geslaagd" lijken. Regressietest voor een kritieke
audit-fix: vóór de fix werd het per-paragraaf resultaat van _sjb_vervang_in_par
nergens opgevangen, dus zo'n mismatch gaf geen enkel signaal — de originele
voorbeeldtekst bleef permanent op elke toekomstige offerte/factuur staan, terwijl
de UI "Sjabloon opgeslagen!" toonde.

Bouwt een echt .docx-document in-memory met python-docx en roept de ECHTE
sjabloon_templatiseer() aan (AST-extractie) — geen Streamlit-runtime nodig.
"""
import copy
import io

from docx import Document

from _schildertool_extract import extract

NAMES = (
    "_sjb_vervang_in_par", "_sjb_iter_paragrafen", "_sjb_zet_celtekst",
    "_sjb_tag_rij", "SJABLOON_VELDEN", "sjabloon_templatiseer",
)


def _maak_ns():
    return extract(*NAMES, extra_globals={"_sjb_io": io, "_sjb_copy": copy})


def _voorbeelddocument(regels):
    doc = Document()
    for regel in regels:
        doc.add_paragraph(regel)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _alle_tekst(docx_bytes):
    doc = Document(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


def test_exact_matchend_veld_wordt_vervangen_en_niet_in_niet_vervangen():
    ns = _maak_ns()
    bron = _voorbeelddocument(["Klant: Jan Jansen", "Adres: Kerkstraat 12"])
    mapping = {"velden": {"klantnaam": "Jan Jansen"}, "tabel": None}

    nieuw, niet_vervangen = ns["sjabloon_templatiseer"](bron, mapping)

    assert niet_vervangen == []
    assert "{{ klantnaam }}" in _alle_tekst(nieuw)
    assert "Jan Jansen" not in _alle_tekst(nieuw)


def test_niet_matchende_tekst_landt_in_niet_vervangen_met_leesbaar_label():
    """Het kernscenario: een handmatige correctie in het bevestigingsformulier die
    niet meer letterlijk in het document voorkomt (bv. een typfout gecorrigeerd,
    of net een net ander format)."""
    ns = _maak_ns()
    bron = _voorbeelddocument(["Klant: Jan Jansen"])
    # 'Janssen' (dubbele s) komt NERGENS letterlijk voor in het document.
    mapping = {"velden": {"klantnaam": "Jan Janssen"}, "tabel": None}

    nieuw, niet_vervangen = ns["sjabloon_templatiseer"](bron, mapping)

    assert niet_vervangen == ["Klantnaam"], (
        "BUG NIET GEFIXT: een niet-matchende tekst moet expliciet gerapporteerd "
        "worden, niet stil verdwijnen"
    )
    # De oorspronkelijke voorbeeldtekst blijft (bewust) onaangeroerd staan — dat is
    # precies waarom de aanroeper dit aan de gebruiker moet melden.
    assert "Jan Jansen" in _alle_tekst(nieuw)
    assert "{{ klantnaam }}" not in _alle_tekst(nieuw)


def test_mix_van_matchend_en_niet_matchend_veld():
    ns = _maak_ns()
    bron = _voorbeelddocument(["Klant: Jan Jansen", "Project: Verbouwing keuken"])
    mapping = {"velden": {
        "klantnaam": "Jan Jansen",              # matcht letterlijk
        "projectnaam": "Verbouwing badkamer",   # matcht NIET (afwijkende tekst)
    }, "tabel": None}

    nieuw, niet_vervangen = ns["sjabloon_templatiseer"](bron, mapping)

    assert niet_vervangen == ["Projectnaam"]
    tekst = _alle_tekst(nieuw)
    assert "{{ klantnaam }}" in tekst
    assert "{{ projectnaam }}" not in tekst
    assert "Verbouwing keuken" in tekst   # onaangeroerd, want niet gematcht


def test_bekend_veld_krijgt_het_leesbare_label_in_niet_vervangen():
    ns = _maak_ns()
    bron = _voorbeelddocument(["niets relevants hier"])
    mapping = {"velden": {"totaal_incl_btw": "€ 999,00"}, "tabel": None}

    _, niet_vervangen = ns["sjabloon_templatiseer"](bron, mapping)

    # SJABLOON_VELDEN kent dit veld een leesbaar Nederlands label toe -> dat label
    # hoort de gebruiker te zien, niet de interne sleutelnaam.
    assert niet_vervangen == ["Totaal incl. btw"]


def test_onbekend_veld_valt_terug_op_de_veldsleutel_zelf_als_label():
    ns = _maak_ns()
    bron = _voorbeelddocument(["Lorem ipsum dolor sit amet"])
    mapping = {"velden": {"dit_bestaat_niet_in_sjabloon_velden": "xyzzy-nergens-in-doc"}, "tabel": None}

    _, niet_vervangen = ns["sjabloon_templatiseer"](bron, mapping)

    # Onbekende veldsleutel -> SJABLOON_VELDEN.get(veld, veld) valt terug op de
    # ruwe sleutel. Belangrijkste is dat het veld gemeld wordt, niet stil verdwijnt.
    assert niet_vervangen == ["dit_bestaat_niet_in_sjabloon_velden"]


def test_leeg_veld_wordt_genegeerd_niet_als_mismatch_gemeld():
    """Een leeg/whitespace-only veld hoort niet als 'niet gevonden' te tellen —
    dat is gewoon een veld dat de gebruiker bewust leeg heeft gelaten."""
    ns = _maak_ns()
    bron = _voorbeelddocument(["Klant: Jan Jansen"])
    mapping = {"velden": {"klantnaam": "Jan Jansen", "website": "   "}, "tabel": None}

    _, niet_vervangen = ns["sjabloon_templatiseer"](bron, mapping)

    assert niet_vervangen == []
